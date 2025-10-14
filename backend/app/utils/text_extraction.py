"""
Text Extraction Utilities
Extract text from various file formats (PDF, DOCX, TXT)
"""

import io
import re
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
import structlog

logger = structlog.get_logger(__name__)


class TextExtractionError(Exception):
    """Custom exception for text extraction errors"""
    pass


def normalize_text(text: str) -> str:
    """
    Normalize extracted text by removing excessive whitespace and newlines.

    This fixes issues with PDF extraction where each word is on a separate line.

    Args:
        text: Raw extracted text

    Returns:
        Normalized text with proper spacing
    """
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)

    # First, collapse lines with words into paragraphs
    # Only treat 2+ consecutive empty lines as paragraph breaks
    lines = text.split('\n')
    all_words = []
    consecutive_empty = 0

    for line in lines:
        stripped = line.strip()

        if not stripped:
            consecutive_empty += 1
            # After 2+ empty lines, insert a paragraph marker
            if consecutive_empty >= 2 and all_words and all_words[-1] != '\n\n':
                all_words.append('\n\n')
        else:
            consecutive_empty = 0
            # Add all words from this line
            words = stripped.split()
            all_words.extend(words)

    # Join all words with spaces, preserving paragraph breaks
    normalized = ' '.join(all_words)

    # Clean up punctuation spacing
    # Remove spaces before punctuation
    normalized = re.sub(r'\s+([.,;:!?)\]])', r'\1', normalized)

    # Remove spaces after opening brackets/quotes
    normalized = re.sub(r'([\(\[])\s+', r'\1', normalized)

    # Ensure space after punctuation (but not for decimals or abbreviations)
    normalized = re.sub(r'([.,;:!?])([A-Za-z])', r'\1 \2', normalized)

    # Fix common patterns
    # Fix markdown bold that got split: ** word -> **word
    normalized = re.sub(r'\*\*\s+', '**', normalized)
    normalized = re.sub(r'\s+\*\*', '**', normalized)

    # Fix hyphenated bullet points: - word -> - word
    normalized = re.sub(r'(?<=\n\n)-\s+', '- ', normalized)
    normalized = re.sub(r'^-\s+', '- ', normalized)

    # Fix > quote markers
    normalized = re.sub(r'>\s+', '> ', normalized)

    # Remove multiple consecutive spaces (but not newlines)
    normalized = re.sub(r'(?<!\n) {2,}(?!\n)', ' ', normalized)

    # Clean up excessive blank lines
    normalized = re.sub(r'\n\n\s+\n\n', '\n\n', normalized)
    normalized = re.sub(r'\n{3,}', '\n\n', normalized)

    return normalized.strip()


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from TXT file

    Args:
        file_content: File content as bytes

    Returns:
        Extracted text

    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        # Try UTF-8 first, then fallback to latin-1
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text = file_content.decode('latin-1')

        logger.info("Text extracted from TXT", length=len(text))
        return text.strip()

    except Exception as e:
        logger.error("Failed to extract text from TXT", error=str(e))
        raise TextExtractionError(f"Failed to extract text from TXT file: {str(e)}")


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file using PyPDF2

    Args:
        file_content: File content as bytes

    Returns:
        Extracted text

    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        if len(reader.pages) == 0:
            raise TextExtractionError("PDF file has no pages")

        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                else:
                    logger.warning(f"No text found on page {page_num}")
            except Exception as e:
                logger.error(f"Failed to extract text from page {page_num}", error=str(e))
                continue

        if not text_parts:
            raise TextExtractionError("No text could be extracted from PDF")

        text = "\n\n".join(text_parts)

        # Normalize text to fix excessive newlines from PDF extraction
        normalized_text = normalize_text(text)

        logger.info(
            "Text extracted from PDF",
            pages=len(reader.pages),
            original_length=len(text),
            normalized_length=len(normalized_text)
        )
        return normalized_text

    except TextExtractionError:
        raise
    except Exception as e:
        logger.error("Failed to extract text from PDF", error=str(e))
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file using python-docx

    Args:
        file_content: File content as bytes

    Returns:
        Extracted text

    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        # Extract text from all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        if not paragraphs:
            # Try to extract from tables if no paragraphs
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

        if not paragraphs:
            raise TextExtractionError("No text could be extracted from DOCX")

        text = "\n\n".join(paragraphs)
        logger.info(
            "Text extracted from DOCX",
            paragraphs=len(paragraphs),
            length=len(text)
        )
        return text.strip()

    except TextExtractionError:
        raise
    except Exception as e:
        logger.error("Failed to extract text from DOCX", error=str(e))
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_content: bytes, file_type: str) -> str:
    """
    Extract text from file based on type

    Args:
        file_content: File content as bytes
        file_type: Type of file ('txt', 'pdf', 'docx')

    Returns:
        Extracted text

    Raises:
        TextExtractionError: If extraction fails
        ValueError: If file type is not supported
    """
    file_type_lower = file_type.lower()

    extractors = {
        'txt': extract_text_from_txt,
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
    }

    extractor = extractors.get(file_type_lower)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor(file_content)


def validate_extracted_text(text: str, min_length: int = 10) -> bool:
    """
    Validate that extracted text meets minimum requirements

    Args:
        text: Extracted text
        min_length: Minimum required text length

    Returns:
        True if valid, False otherwise
    """
    if not text or not text.strip():
        return False

    if len(text.strip()) < min_length:
        logger.warning(
            "Extracted text too short",
            length=len(text.strip()),
            min_required=min_length
        )
        return False

    return True
